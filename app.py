import email

from flask import Flask, flash, render_template, request
from flask import session, redirect, url_for
from werkzeug.utils import secure_filename
import os
import fitz
from utils import extract_entities
import re
import requests
import html
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from flask import send_file
from reportlab.platypus import Image
import sqlite3
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from datetime import datetime
from datetime import datetime
from authlib.integrations.flask_client import OAuth
import uuid
from werkzeug.utils import secure_filename
from flask_mail import Mail, Message


UPLOAD_FOLDER = 'static/uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key')


app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.environ.get('MAIL_USERNAME')

mail = Mail(app)
if not app.config['MAIL_USERNAME'] or not app.config['MAIL_PASSWORD']:
    print("⚠️ Email config missing")

from itsdangerous import URLSafeTimedSerializer

s = URLSafeTimedSerializer(app.config['SECRET_KEY'])
# =========================================
# 🌍 GLOBAL USER FOR ALL TEMPLATES
# =========================================
@app.context_processor
def inject_user():
    if 'user' in session:
        conn = sqlite3.connect('users.db')
        cursor = conn.cursor()

        cursor.execute("""
            SELECT firstname, lastname, profile_image
            FROM users
            WHERE username=?
        """, (session['user'],))

        data = cursor.fetchone()
        conn.close()

        if data:
            return {
                "user": {
                    "firstname": data[0],
                    "lastname": data[1],
                    "profile_image": data[2]
                }
            }

    return {"user": None}

oauth = OAuth(app)

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email')

        conn = sqlite3.connect('users.db')
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM users WHERE username=?", (email,))
        user = cursor.fetchone()

        conn.close()

        if user:
            token = s.dumps(email, salt='reset-password')
            reset_link = url_for('reset_password', token=token, _external=True)

            msg = Message(
                subject="Password Reset Request",
                recipients=[email]
            )

            msg.body = f"""Hi {user[1]},

You requested a password reset. Click the link below to reset your password:

{reset_link}

If you didn't request this, please ignore this email.

Best,
AI Resume Analyzer Team
"""

            # ✅ SAFE EMAIL SEND (IMPORTANT FIX)
            try:
                mail.send(msg)
            except Exception as e:
                print("Email error:", e)
                return render_template(
                    'forgot_password.html',
                    error="Failed to send email. Please try again later."
                )

            return render_template(
                'forgot_password.html',
                message="Reset link sent to your email!"
            )
        else:
            return render_template(
                'forgot_password.html',
                error="Email not found"
            )

    return render_template('forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    try:
        email = s.loads(token, salt='reset-password', max_age=300)
    except:
        return render_template('reset_password.html', error="Invalid or expired link", token=token)

    if request.method == 'POST':
        new_password = request.form.get('newPassword')
        confirm_password = request.form.get('confirmPassword')

        if not new_password or not confirm_password:
            return render_template('reset_password.html', error="All fields are required", token=token)

        if new_password != confirm_password:
            return render_template('reset_password.html', error="Passwords do not match", token=token)

        conn = sqlite3.connect('users.db')
        cursor = conn.cursor()

        cursor.execute(
            "UPDATE users SET password=? WHERE username=?",
            (new_password, email)
        )
        conn.commit()
        conn.close()

        return render_template('reset_password.html', message="Password updated successfully", token=token)

    return render_template('reset_password.html', token=token)

@app.route('/callback/google')
def google_callback():
    token = google.authorize_access_token()
    user_info = google.get('userinfo').json()

    email = user_info['email']
    name = user_info['name']
    picture = user_info.get('picture')   # 🔥 GET IMAGE

    filename = None

    # =========================
    # 📸 DOWNLOAD IMAGE
    # =========================
    if picture:
        response = requests.get(picture)

        if response.status_code == 200:
            filename = str(uuid.uuid4()) + ".png"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

            with open(filepath, 'wb') as f:
                f.write(response.content)

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM users WHERE username=?", (email,))
    user = cursor.fetchone()

    if not user:
        cursor.execute("""
            INSERT INTO users (firstname, lastname, username, profile_image)
            VALUES (?, ?, ?, ?)
        """, (name.split()[0], name.split()[-1], email, filename))
    else:
        # update image if user already exists
        cursor.execute("""
            UPDATE users SET profile_image=?
            WHERE username=?
        """, (filename, email))

    conn.commit()
    conn.close()

    session['user'] = email
    session['firstname'] = name.split()[0]
    session['lastname'] = name.split()[-1]
    session['profile_image'] = filename   # 🔥 IMPORTANT

    return redirect(url_for('dashboard'))




google = oauth.register(
    name='google',
    client_id=os.environ.get("GOOGLE_CLIENT_ID"),
    client_secret=os.environ.get("GOOGLE_CLIENT_SECRET"),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email profile'
    }
)

# ✅ ADD HERE
if not os.environ.get("GOOGLE_CLIENT_ID") or not os.environ.get("GOOGLE_CLIENT_SECRET"):
    print("⚠️ Google OAuth config missing")

@app.route('/login/google')
def login_google():
    return google.authorize_redirect(url_for('authorize_google', _external=True))

@app.route('/authorize/google')
def authorize_google():
    token = google.authorize_access_token()
    user = google.get('https://www.googleapis.com/oauth2/v3/userinfo').json()

    # Store user in session
    session['user'] = user['email']

    return redirect(url_for('dashboard'))


@app.route('/update_profile', methods=['POST'])
def update_profile():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    # =========================
    # 📸 HANDLE IMAGE UPLOAD
    # =========================
    file = request.files.get('profile_image')
    filename = None

    if file and file.filename != '':
        filename = str(uuid.uuid4()) + "_" + secure_filename(file.filename)
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

    # =========================
    # 🧠 GET OLD IMAGE (SAFE)
    # =========================
    cursor.execute("SELECT profile_image FROM users WHERE username=?", (session['user'],))
    result = cursor.fetchone()

    old_image = result[0] if result and result[0] else None

    if not filename:
        filename = old_image

    # =========================
    # 📝 GET FORM DATA
    # =========================
    firstname = request.form.get('firstname')
    lastname = request.form.get('lastname')
    email = request.form.get('email')
    phone = request.form.get('phone')
    location = request.form.get('location')
    title = request.form.get('title')
    company = request.form.get('company')
    years_of_experience = request.form.get('experience')
    industry = request.form.get('industry')
    summary = request.form.get('summary')
    skills = request.form.get('skills')
    linkedin = request.form.get('linkedin')
    github = request.form.get('github')

    # =========================
    # ✅ UPDATE USER
    # =========================
    cursor.execute("""
UPDATE users
SET firstname=?, lastname=?, username=?, phone=?, location=?, title=?, company=?, 
skills=?, linkedin=?, github=?, industry=?, summary=?, years_of_experience=?, profile_image=?
WHERE username=?
""", (
    firstname, lastname, email,
    phone, location, title, company,
    skills, linkedin, github,
    industry, summary, years_of_experience,
    filename,
    session['user']
))

    conn.commit()
    conn.close()

    # =========================
    # 🔄 UPDATE SESSION
    # =========================
    session['user'] = email
    session['firstname'] = firstname
    session['lastname'] = lastname
    session['profile_image'] = filename   # 🔥 IMPORTANT ADD

    return redirect(url_for('profile'))


@app.route('/delete_account', methods=['POST'])
def delete_account():
    if 'user' not in session:
        return redirect(url_for('login'))

    username = session['user']

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    # =========================
    # 📸 DELETE PROFILE IMAGE (PUT HERE ✅)
    # =========================
    cursor.execute("SELECT profile_image FROM users WHERE username=?", (username,))
    img = cursor.fetchone()

    if img and img[0]:
        path = os.path.join(app.config['UPLOAD_FOLDER'], img[0])
        if os.path.exists(path):
            os.remove(path)

    # =========================
    # 🗑 DELETE DATA FROM TABLES
    # =========================
    cursor.execute("DELETE FROM users WHERE username=?", (username,))
    cursor.execute("DELETE FROM analysis_history WHERE username=?", (username,))
    cursor.execute("DELETE FROM skills_history WHERE username=?", (username,))

    conn.commit()
    conn.close()

    # =========================
    # 🔄 CLEAR SESSION
    # =========================
    session.clear()

    return redirect(url_for('login'))

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
def init_db():
    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()
    
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS analysis_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT,
        job_title TEXT,
        score REAL,
        date TEXT
    )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            firstname TEXT,
            lastname TEXT,
            username TEXT UNIQUE,
            password TEXT
        )
    ''')
    
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS skills_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT,
        skill TEXT,
        score REAL
    )
    ''')

    conn.commit()
    conn.close()

def extract_text_pymupdf(pdf_path):
    text = ""
    doc = fitz.open(pdf_path)
    for page in doc:
        text += page.get_text()
    doc.close()
    return text.strip()





# 🔥 Highlight function (KEEP ABOVE analyze)
def highlight_text(text, matched, missing):
    highlighted = html.escape(text)

    # Matched skills → green
    for skill in matched:
        highlighted = re.sub(
            rf"\b({re.escape(skill)})\b",
            r'<span style="background-color: #bbf7d0">\1</span>',
            highlighted,
            flags=re.IGNORECASE
        )

    # Missing skills → red
    for skill in missing:
        highlighted = re.sub(
            rf"\b({re.escape(skill)})\b",
            r'<span style="background-color: #fecaca">\1</span>',
            highlighted,
            flags=re.IGNORECASE
        )

    return highlighted

def recommend_jobs(resume_skills):
    job_scores = {}

    for job, skills in JOB_ROLES.items():
        score = len(set(skills) & resume_skills)
        if score > 0:
            job_scores[job] = score

    # Sort by best match
    sorted_jobs = sorted(job_scores, key=job_scores.get, reverse=True)

    return sorted_jobs[:5]


@app.route('/')
def index():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    cursor.execute("""
        SELECT firstname, lastname, profile_image
        FROM users
        WHERE username=?
    """, (session['user'],))

    data = cursor.fetchone()
    conn.close()

    user = None
    if data:
        user = {
            "firstname": data[0],
            "lastname": data[1],
            "profile_image": data[2]
        }

    return render_template('index.html', user=user)

@app.route('/profile')
def profile():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    # =========================
    # 👤 USER DATA
    # =========================
    cursor.execute("""
    SELECT firstname, lastname, username, phone, location, title, company,
           skills, linkedin, github, industry, summary, years_of_experience, profile_image
    FROM users
    WHERE username=?
    """, (session['user'],))

    user_data = cursor.fetchone()

    # =========================
    # 📊 STATS
    # =========================
    cursor.execute("""
        SELECT COUNT(*), AVG(score)
        FROM analysis_history
        WHERE username=?
    """, (session['user'],))

    stats = cursor.fetchone()

    total_analysis = stats[0] or 0
    avg_score = round(stats[1], 2) if stats[1] else 0

    # ✅ same as total analysis
    resumes_uploaded = total_analysis

    # ✅ temporary (you can make DB later)
    profile_views = 0

    conn.close()

    if user_data:
        user = {
            "firstname": user_data[0],
            "lastname": user_data[1],
            "email": user_data[2],
            "phone": user_data[3],
            "location": user_data[4],
            "title": user_data[5],
            "company": user_data[6],
            "skills": user_data[7],
            "linkedin": user_data[8],
            "github": user_data[9],
            "industry": user_data[10],
            "summary": user_data[11],
            "experience": user_data[12], 
            "profile_image": user_data[13],

            # 🔥 FIXED KEYS
            "total_analysis": total_analysis,
            "resumes_uploaded": resumes_uploaded,   # ✅ FIXED NAME
            "avg_score": avg_score,
            "profile_views": profile_views          # ✅ ADDED
        }
    else:
        user = {}

    return render_template('profile.html', user=user)



# @app.route('/templates')
# def templates():
#     if 'user' not in session:
#         return redirect(url_for('login'))

#     templates_data = [
#         {
#             "name": "Classic Professional",
#             "image": "temp-1.png",
#             "category": "professional",
#             "description": "Clean, ATS-friendly layout",
#             "featured": True
#         },
#         {
#             "name": "Modern Minimal",
#             "image": "temp-2.png",
#             "category": "modern",
#             "description": "Sleek two-column design",
#             "featured": False
#         },
#         {
#             "name": "Executive Bold",
#             "image": "temp-3.png",
#             "category": "executive",
#             "description": "Premium executive style",
#             "featured": False
#         },
#         {
#             "name": "Creative Portfolio",
#             "image": "temp-2.png",
#             "category": "creative",
#             "description": "Showcase creative work",
#             "featured": True
#         },
#         {
#             "name": "Technical Developer",
#             "image": "temp-5.png",
#             "category": "professional",
#             "description": "Perfect for tech roles",
#             "featured": False
#         },
#         {
#             "name": "Academic Scholar",
#             "image": "temp-6.png",
#             "category": "executive",
#             "description": "Academic-focused design",
#             "featured": False
#         }
#     ]

#     return render_template("resume-temp.html", templates=templates_data)

@app.route('/history')
def history():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    # 🔥 FETCH HISTORY
    cursor.execute("""
        SELECT job_title, score, date
        FROM analysis_history
        WHERE username=?
        ORDER BY id DESC
    """, (session['user'],))

    history_data = cursor.fetchall()

    # =========================
    # 📊 CALCULATE STATS
    # =========================
    total_analyses = len(history_data)

    avg_score = round(
        sum([row[1] for row in history_data]) / total_analyses, 2
    ) if total_analyses > 0 else 0

    best_score = max([row[1] for row in history_data]) if total_analyses > 0 else 0

    # This month count
    from datetime import datetime
    current_month = datetime.now().strftime("%Y-%m")

    this_month_count = sum(
        1 for row in history_data if row[2].startswith(current_month)
    )

    conn.close()

    return render_template(
        'history.html',
        history=history_data,
        total_analyses=total_analyses,
        avg_score=avg_score,
        best_score=best_score,
        this_month_count=this_month_count
    )
@app.route('/dashboard')
def dashboard():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    cursor.execute('''
        SELECT job_title, score, date 
        FROM analysis_history 
        WHERE username=? 
        ORDER BY id DESC
    ''', (session['user'],))

    history = cursor.fetchall()

    total_analysis = len(history)

    avg_score = round(
        sum([row[1] for row in history]) / total_analysis, 2
    ) if total_analysis > 0 else 0

    resumes_uploaded = total_analysis

    cursor.execute('''
        SELECT COUNT(DISTINCT job_title)
        FROM analysis_history
        WHERE username=?
    ''', (session['user'],))

    job_desc_count = cursor.fetchone()[0]

    cursor.execute('''
        SELECT skill, AVG(score) as avg_score
        FROM skills_history
        WHERE username=?
        GROUP BY skill
        ORDER BY avg_score DESC
        LIMIT 5
    ''', (session['user'],))

    top_skills_data = cursor.fetchall()

    top_skills = []
    for skill, score in top_skills_data:
        top_skills.append({
            "name": skill.capitalize(),
            "score": round(score, 2)
        })

    conn.close()

    return render_template(
        'dashboard.html',
        history=history,
        total_analysis=total_analysis,
        avg_score=avg_score,
        resumes_uploaded=resumes_uploaded,
        job_desc_count=job_desc_count,
        top_skills=top_skills
    )
    
@app.route('/chatbot')
def chatbot_page():
    if 'user' not in session:
        return redirect(url_for('login'))

    data = session.get('report_data', {})

    return render_template(
        'chatbot.html',
        matched_skills=data.get('matched_skills', []),
        missing_skills=data.get('missing_skills', []),
        match_score=data.get('match_score', 0),
        suggestions=data.get('suggestions', [])
    )

# 🔐 LOGIN ROUTE
@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        conn = sqlite3.connect('users.db')
        cursor = conn.cursor()

        cursor.execute(
            "SELECT firstname, lastname, password FROM users WHERE username=?",
            (username,)
        )
        user = cursor.fetchone()

        conn.close()

        # ✅ SINGLE CONDITION FOR BOTH CASES
        if user and user[2] == password:
            session['user'] = username
            session['firstname'] = user[0]
            session['lastname'] = user[1]

            return redirect(url_for('dashboard'))
        else:
            error = "Invalid username or password"

    return render_template('login.html', error=error)


# 📝 SIGNUP ROUTE
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        firstname = request.form.get('firstname')   # ✅ NEW
        lastname = request.form.get('lastname')     # ✅ NEW
        username = request.form.get('username')
        password = request.form.get('password')

        conn = sqlite3.connect('users.db')
        cursor = conn.cursor()

        try:
            cursor.execute(
                "INSERT INTO users (firstname, lastname, username, password) VALUES (?, ?, ?, ?)",
                (firstname, lastname, username, password)
            )
            conn.commit()
        except:
            return "User already exists"

        conn.close()
        return redirect(url_for('login'))

    return render_template('signup.html')

# 🚪 LOGOUT ROUTE
@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('login'))

JOB_ROLES = {
    "Python Developer": ["python", "flask", "django", "api"],
    "Java Developer": ["java", "spring", "spring boot"],
    "Frontend Developer": ["html", "css", "javascript", "react"],
    "Data Scientist": ["machine learning", "numpy", "pandas", "nlp"],
    "Backend Developer": ["api", "sql", "database"],
    "DevOps Engineer": ["docker", "aws", "ci", "cd"]
}



from datetime import datetime
@app.route('/analyze', methods=['POST'])
def analyze():
    resume_file = request.files['resume']
    job_desc = request.form['job_description']

    if not resume_file or not job_desc:
        return "Resume file or job description missing.", 400

    # ================================
    # 📄 SAVE FILE
    # ================================
    filename = secure_filename(resume_file.filename)
    save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    resume_file.save(save_path)

    # ================================
    # 📄 EXTRACT TEXT
    # ================================
    resume_text = extract_text_pymupdf(save_path)

    # ================================
    # 🔥 SKILL EXTRACTION
    # ================================
    resume_entities = extract_entities(resume_text)
    resume_skills = set(w.lower() for w in resume_entities.get('SKILL', []))

    # ================================
    # 🔥 JOB RECOMMENDATION
    # ================================
    recommended_jobs = recommend_jobs(resume_skills)

    # ================================
    # 🔥 NEW MATCHING LOGIC (FIXED)
    # ================================
    resume_text_lower = resume_text.lower()
    job_desc_lower = job_desc.lower()

    jd_skills = set()

    for category, skills in SKILL_CATEGORIES.items():
        for skill in skills:
            if skill in resume_text_lower:
                resume_skills.add(skill)
            if skill in job_desc_lower:
                jd_skills.add(skill)

    matched_skills = list(resume_skills & jd_skills)
    missing_skills = list(jd_skills - resume_skills)

    # ✅ SCORE (IMPORTANT FIX)
    if jd_skills:
        match_score = round((len(matched_skills) / len(jd_skills)) * 100, 2)
    else:
        match_score = 0

    # ================================
    # 🔥 HIGHLIGHT RESUME
    # ================================
    highlighted_resume = highlight_text(resume_text, matched_skills, missing_skills)

    # ================================
    # 💡 SMART SUGGESTIONS
    # ================================
    suggestions = []

    if match_score < 60:
        suggestions.append("Improve your resume by aligning it with the job description.")

    if missing_skills:
        suggestions.append("Missing important skills: " + ", ".join(missing_skills[:5]))

    if len(resume_text.split()) < 150:
        suggestions.append("Your resume is too short. Add more detailed content.")

    if not any(h in resume_text.lower() for h in ["experience", "project", "internship"]):
        suggestions.append("Add an Experience or Projects section.")

    if len(resume_skills) < 5:
        suggestions.append("Include more technical skills.")

    # ================================
    # 🔥 SAVE TO DATABASE
    # ================================
    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    job_title = job_desc.split('\n')[0][:50]

    cursor.execute('''
        INSERT INTO analysis_history (username, job_title, score, date)
        VALUES (?, ?, ?, ?)
    ''', (
        session['user'],
        job_title,
        match_score,
        datetime.now().strftime("%Y-%m-%d")
    ))

    for skill in matched_skills:
        skill_score = round(match_score * (0.8 + len(skill)/20), 2)
        skill_score = min(skill_score, 100)

        cursor.execute('''
            INSERT INTO skills_history (username, skill, score)
            VALUES (?, ?, ?)
        ''', (
            session['user'],
            skill,
            skill_score
        ))

    conn.commit()
    conn.close()

    # ================================
    # 🔥 TOP SKILLS
    # ================================
    top_skills = []

    for skill in matched_skills[:5]:
        score = round(match_score * (0.8 + len(skill)/20), 2)
        score = min(score, 100)

        top_skills.append({
            "name": skill.capitalize(),
            "score": score
        })

    # Save session
    session['report_data'] = {
        "match_score": match_score,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "suggestions": suggestions,
        "top_skills": top_skills
    }

    return render_template(
        'result.html',
        match_score=match_score,
        matched_skills=matched_skills,
        missing_skills=missing_skills,
        suggestions=suggestions,
        highlighted_resume=highlighted_resume,
        recommended_jobs=recommended_jobs,
        top_skills=top_skills
    )
    
from flask import jsonify

@app.route('/chat', methods=['POST'])
def chat():
    data = request.json

    user_question = data.get("message", "").lower()
    matched_skills = data.get("matched_skills", [])
    missing_skills = data.get("missing_skills", [])
    score = data.get("score", 0)
    suggestions = data.get("suggestions", [])

    reply = ""

    # 🔥 Smart responses
    if "improve" in user_question:
        reply = f"To improve your resume, you should focus on adding these skills: {', '.join(missing_skills[:5])}."

    elif "skills" in user_question:
        reply = f"You already have strong skills like {', '.join(matched_skills[:5])}. Consider learning {', '.join(missing_skills[:3])}."

    elif "missing" in user_question:
        reply = f"You are missing important skills such as: {', '.join(missing_skills[:5])}."

    elif "score" in user_question:
        reply = f"Your resume match score is {score}%. Improve it by aligning your skills with the job description."

    elif "job" in user_question:
        reply = "Based on your skills, you can apply for roles like Backend Developer, Data Analyst, or Python Developer."

    elif "project" in user_question:
        reply = "Add strong projects related to your domain. Example: Resume Analyzer, Web App, or ML Project."

    elif "hello" in user_question or "hi" in user_question:
        reply = "Hello 👋 I am your Resume Assistant. Ask me anything about your resume."

    else:
        # fallback smart answer
        reply = "Based on your resume, focus on improving skills, adding projects, and aligning with job requirements."

    return jsonify({"reply": reply})
    
    

    

    
@app.route('/download_pdf', methods=['POST'])
def download_pdf():

    data = session.get('report_data')

    if not data:
        return "No report data found"

    match_score = data['match_score']
    matched_list = data['matched_skills']
    missing_list = data['missing_skills']
    suggestions = data['suggestions']
    top_skills = data['top_skills']

    file_path = "report.pdf"

    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    elements = []

    # ================================
    # 🔥 HEADER (PROFESSIONAL)
    # ================================
    header = Table([
        ["AI Resume Analyzer Report"]
    ], colWidths=[450])

    header.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 0), (-1, -1), 18),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12)
    ]))

    elements.append(header)
    elements.append(Spacer(1, 15))

    # ================================
    # 📊 SCORE BOX
    # ================================
    score_table = Table([
        ["Match Score", f"{match_score}%"]
    ], colWidths=[200, 200])

    score_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 0), (-1, -1), 14),
        ('BOX', (0, 0), (-1, -1), 1, colors.black)
    ]))

    elements.append(score_table)
    elements.append(Spacer(1, 20))

    # ================================
    # ✅ MATCHED SKILLS
    # ================================
    elements.append(Paragraph("<b>Matched Skills</b>", styles['Heading2']))
    for skill in matched_list:
        elements.append(Paragraph(f"✔ {skill}", styles['Normal']))
    elements.append(Spacer(1, 12))

    # ================================
    # ❌ MISSING SKILLS
    # ================================
    elements.append(Paragraph("<b>Missing Skills</b>", styles['Heading2']))
    for skill in missing_list:
        elements.append(Paragraph(f"✘ {skill}", styles['Normal']))
    elements.append(Spacer(1, 12))

    # ================================
    # ⭐ TOP SKILLS (TABLE)
    # ================================
    elements.append(Paragraph("<b>Top Skills</b>", styles['Heading2']))

    skill_data = [["Skill", "Score"]]

    for skill in top_skills:
        skill_data.append([skill['name'], f"{skill['score']}%"])

    skill_table = Table(skill_data)

    skill_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))

    elements.append(skill_table)
    elements.append(Spacer(1, 15))

    # ================================
    # 💡 SUGGESTIONS
    # ================================
    elements.append(Paragraph("<b>Suggestions</b>", styles['Heading2']))

    for s in suggestions:
        elements.append(Paragraph(f"• {s}", styles['Normal']))

    elements.append(Spacer(1, 20))

    # ================================
    # 📄 FOOTER
    # ================================
    elements.append(Paragraph(
        "<i>Generated by AI Resume Analyzer</i>",
        styles['Normal']
    ))

    # BUILD PDF
    doc.build(elements)

    return send_file(file_path, as_attachment=True)

@app.route('/create_resume')
def create_resume():
    return render_template('create_resume.html')

@app.route('/edit-profile')
def edit_profile():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    cursor.execute("""
SELECT firstname, lastname, username, phone, location, title, company,
       skills, linkedin, github, profile_image,
       years_of_experience, summary
FROM users
WHERE username=?
""", (session['user'],))

    user_data = cursor.fetchone()
    conn.close()

    if user_data:
        user = {
    "firstname": user_data[0],
    "lastname": user_data[1],
    "email": user_data[2],  # ok if username = email
    "phone": user_data[3],
    "location": user_data[4],
    "title": user_data[5],
    "company": user_data[6],
    "skills": user_data[7],
    "linkedin": user_data[8],
    "github": user_data[9],
    "profile_image": user_data[10],
    "years_of_experience": user_data[11],
    "summary": user_data[12]
}
    else:
        user = {}

    return render_template('edit-profile.html', user=user)



def upgrade_db():
    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()

    columns = [
        "phone", "location", "title", "company",
        "skills", "linkedin", "github",
        "industry", "summary", "years_of_experience",
        "profile_image"   # ✅ ADD THIS
    ]

    for col in columns:
        try:
            cursor.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT")
        except:
            pass  # already exists

    conn.commit()
    conn.close()

from flask import render_template, request, send_file
import os



from docx import Document

@app.route('/download_docx', methods=['POST'])
def download_docx():

    data = request.form

    doc = Document()

    # 🔥 NAME
    doc.add_heading(f"{data.get('firstname')} {data.get('lastname')}", 0)

    # ROLE
    doc.add_paragraph(data.get('role'))

    # CONTACT
    doc.add_paragraph(f"{data.get('email')} | {data.get('phone')}")

    # OBJECTIVE
    doc.add_heading("Career Objective", level=1)
    doc.add_paragraph(data.get('objective'))

    # SKILLS
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(data.get('skills'))

    # EDUCATION
    doc.add_heading("Education", level=1)
    doc.add_paragraph(data.get('education'))

    # PROJECTS
    doc.add_heading("Projects", level=1)
    doc.add_paragraph(data.get('projects'))

    file_path = "resume.docx"
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

SKILL_CATEGORIES = {
    "Backend": ["java", "spring", "spring boot", "hibernate", "node", "django"],
    "Frontend": ["html", "css", "javascript", "react", "angular"],
    "Database": ["mysql", "sql", "postgresql", "mongodb"],
    "DevOps": ["docker", "kubernetes", "aws", "ci", "cd"],
    "API": ["rest", "api", "microservices"],
    "Tools": ["git", "github", "maven", "gradle"],
}

init_db()
upgrade_db()
if __name__ == '__main__':
    port = int(os.environ.get("PORT", 10000))
    app.run(host='0.0.0.0', port=port)